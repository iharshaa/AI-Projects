from docx import Document
import os

def itinerary_agent(state: dict) -> dict:
    origin = state.get("origin", "")
    location = state.get("location", "")
    doc = Document()

    doc.add_heading(f"{origin} to {location} Trip Itinerary", 0)
    doc.add_paragraph(state.get("itinerary_text", "Itinerary not generated."))

    if not os.path.exists("outputs"):
        os.makedirs("outputs")
    output_path = os.path.join("outputs", f"{origin}_to_{location}_Trip_Itinerary.docx")
    doc.save(output_path)
    state["final_doc_path"] = output_path
    return state